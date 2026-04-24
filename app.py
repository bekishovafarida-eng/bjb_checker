from flask import Flask, render_template, request, redirect, url_for, send_from_directory
from google.cloud import vision
from google import genai
from dotenv import load_dotenv
from docx import Document
from werkzeug.utils import secure_filename
import os
import json
import re

load_dotenv()

app = Flask(__name__)

# ===== ПАПКАЛАР =====
UPLOAD_FOLDER = "uploads"
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# ===== GOOGLE / API =====
google_json = os.getenv("GOOGLE_CREDENTIALS_JSON")
if not google_json:
    raise ValueError("GOOGLE_CREDENTIALS_JSON табылмады")

with open("key.json", "w", encoding="utf-8") as f:
    f.write(google_json)

os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "key.json"

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    raise ValueError("GEMINI_API_KEY табылмады")

vision_client = vision.ImageAnnotatorClient()
gemini_client = genai.Client(api_key=GEMINI_API_KEY)

# ===== JSON PARSER =====
def parse_json_response(text):
    text = text.strip()
    text = text.replace("```json", "").replace("```", "").strip()

    try:
        return json.loads(text)
    except Exception:
        match = re.search(r"\{.*\}", text, re.DOTALL)
        if match:
            return json.loads(match.group(0))
        raise ValueError("JSON оқу кезінде қате шықты")

# ===== OCR =====
def extract_text_from_image(image_file):
    content = image_file.read()
    image = vision.Image(content=content)
    response = vision_client.document_text_detection(image=image)

    if response.full_text_annotation.text:
        return response.full_text_annotation.text
    return "Мәтін табылмады"

# ===== БЖБ ТЕКСЕРУ =====
def check_with_gemini(student_text, answer_key, max_score):
    prompt = f"""
Сен мұғалімнің көмекшісі боласың.
Оқушы жауабын дұрыс жауаппен салыстырып, тексер.

ДҰРЫС ЖАУАП:
{answer_key}

ОҚУШЫ ЖАУАБЫ:
{student_text}

Максималды ұпай: {max_score}

Тек JSON форматында қайтар:
{{
  "score": сан,
  "decision": "Дұрыс/Жартылай дұрыс/Қате",
  "feedback": "Қысқаша қазақша пікір"
}}
"""
    response = gemini_client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt
    )
    return parse_json_response(response.text)

# ===== ЭССЕ ҚҰРУ =====
def generate_essay(topic, grade_level, word_count):
    prompt = f"""
Сен қазақ тілінде эссе жазуға көмектесетін көмекшісің.

Мына параметрлер бойынша толық дайын эссе жаз:
Тақырып: {topic}
Сынып: {grade_level}
Сөз саны: шамамен {word_count} сөз

Талаптар:
1. Эссе қазақ тілінде болсын.
2. Құрылымы анық болсын: кіріспе, негізгі бөлім, қорытынды.
3. Оқушы тіліне жақын, түсінікті болсын.
4. Өте күрделі емес, мектепке сай болсын.
5. Сөз саны шамамен көрсетілген мөлшерге жақын болсын.

Жауапты тек мына JSON форматында қайтар:
{{
  "title": "Эссе тақырыбы",
  "essay": "Толық эссе мәтіні",
  "intro": "Кіріспе қысқаша",
  "main": "Негізгі бөлім қысқаша",
  "conclusion": "Қорытынды қысқаша",
  "estimated_word_count": сан
}}
"""
    response = gemini_client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt
    )
    return parse_json_response(response.text)

# ===== ЭССЕ ТЕКСЕРУ =====
def check_essay_with_gemini(recognized_text, topic="", grade_level="", word_count=""):
    prompt = f"""
Сен эссе тексеретін мұғалім көмекшісісің.

OCR арқылы танылған эссе мәтіні төменде берілген. 
Сен:
1. Мәтінді оқып шық.
2. Қарапайым грамматикалық және орфографиялық қателерді тап.
3. Тақырыпқа сәйкестігін қысқаша бағала.
4. Сөз санына жуық бағалау бер.
5. Жалпы кері байланыс жаз.

Қосымша мәлімет:
Тақырып: {topic}
Сынып: {grade_level}
Күтілетін сөз саны: {word_count}

ТАНЫЛҒАН МӘТІН:
{recognized_text}

Жауапты тек JSON форматында қайтар:
{{
  "recognized_text": "OCR арқылы шыққан мәтін",
  "estimated_word_count": сан,
  "topic_match": "Сәйкес/Жартылай сәйкес/Сәйкес емес",
  "mistakes": [
    "1-қате",
    "2-қате",
    "3-қате"
  ],
  "strengths": [
    "1-артықшылық",
    "2-артықшылық"
  ],
  "feedback": "Қысқаша жалпы пікір",
  "corrected_version": "Мүмкін болса, түзетілген қысқа нұсқа"
}}
"""
    response = gemini_client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt
    )
    return parse_json_response(response.text)

# ===== ТЕСТ ҚҰРУ =====
def generate_test(subject, grade_level, question_count, difficulty):
    prompt = f"""
Сен мұғалімге арналған тест құрастырушы көмекшісің.

Мына параметрлер бойынша қазақ тілінде тест құрастыр:
Пән: {subject}
Сынып: {grade_level}
Сұрақ саны: {question_count}
Деңгейі: {difficulty}

Талаптар:
1. Әр сұраққа 4 нұсқа болсын: A, B, C, D
2. Тест мектепке сай, түсінікті болсын
3. Сұрақтар пәнге және сыныпқа сәйкес болсын
4. Соңында дұрыс жауап кілті берілсін

Жауапты тек JSON форматында қайтар:
{{
  "title": "Тест атауы",
  "questions": [
    {{
      "number": 1,
      "question": "Сұрақ мәтіні",
      "options": {{
        "A": "Нұсқа A",
        "B": "Нұсқа B",
        "C": "Нұсқа C",
        "D": "Нұсқа D"
      }},
      "answer": "A"
    }}
  ],
  "answer_key": {{
    "1": "A",
    "2": "B"
  }}
}}
"""
    response = gemini_client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt
    )
    return parse_json_response(response.text)


# ===== ТЕСТТЕН ЖАУАП КІЛТІН ШЫҒАРУ =====
def extract_answer_key_from_test_text(test_text):
    prompt = f"""
Сен мұғалімнің көмекшісісің.

Төменде тест мәтіні берілген. 
Осы тесттен дұрыс жауап кілтін шығарып бер.

Егер тест мәтінінде жауаптар нақты көрінсе, соларды ал.
Егер сұрақтардың дұрыс жауабын білімге сүйеніп анықтауға болса, анықта.
Әр жауапты тек A/B/C/D форматында бер.

ТЕСТ МӘТІНІ:
{test_text}

Жауапты тек JSON форматында қайтар:
{{
  "answer_key": {{
    "1": "A",
    "2": "B",
    "3": "C"
  }}
}}
"""
    response = gemini_client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt
    )
    return parse_json_response(response.text)


# ===== ОҚУШЫ ЖАУАПТАРЫН PARSE ЖАСАУ =====
def parse_student_answers(answer_text):
    """
    Кіріс мысалы:
    1A 2B 3C
    1-а, 2-б, 3-с
    1)A 2)B 3)C
    """
    normalized = answer_text.upper()
    normalized = normalized.replace("Ә", "A").replace("Б", "B").replace("С", "C").replace("Д", "D")
    normalized = normalized.replace("А", "A").replace("В", "B").replace("С", "C").replace("Д", "D")

    pairs = re.findall(r"(\d+)\s*[\)\.\-:]*\s*([ABCD])", normalized)
    result = {}
    for num, ans in pairs:
        result[num] = ans
    return result


# ===== ТЕСТ ТЕКСЕРУ =====
def check_test_answers(answer_key, student_answers):
    total = len(answer_key)
    correct = 0
    details = []

    for q_num, correct_answer in answer_key.items():
        student_answer = student_answers.get(str(q_num), "-")
        is_correct = student_answer == correct_answer
        if is_correct:
            correct += 1

        details.append({
            "number": q_num,
            "correct_answer": correct_answer,
            "student_answer": student_answer,
            "is_correct": is_correct
        })

    wrong = total - correct
    percent = round((correct / total) * 100, 1) if total > 0 else 0

    return {
        "total": total,
        "correct": correct,
        "wrong": wrong,
        "percent": percent,
        "details": details
    }

# ===== WORD =====
def save_to_word(student_text, result, max_score):
    doc = Document()
    doc.add_heading("БЖБ тексеру нәтижесі", 0)

    doc.add_heading("Оқушы жауабы", level=1)
    doc.add_paragraph(student_text)

    doc.add_heading("Нәтиже", level=1)
    doc.add_paragraph(f"Ұпай: {result['score']} / {max_score}")
    doc.add_paragraph(f"Қорытынды: {result['decision']}")
    doc.add_paragraph(f"Кері байланыс: {result['feedback']}")

    doc.save("result.docx")

# ===== БАСТЫ БЕТ =====
@app.route("/")
def home():
    return render_template("index.html")

# ===== БЖБ =====
@app.route("/bjb-checker", methods=["GET", "POST"])
def bjb():
    result = None
    student_text = ""
    error = ""
    max_score = ""

    if request.method == "POST":
        try:
            images = request.files.getlist("images")
            answer_key = request.form["answer_key"]
            max_score = request.form["max_score"]

            all_texts = []

            for image in images:
                if image and image.filename:
                    text = extract_text_from_image(image)
                    all_texts.append(text)

            student_text = "\n\n----- КЕЛЕСІ БЕТ -----\n\n".join(all_texts)
            result = check_with_gemini(student_text, answer_key, max_score)
            save_to_word(student_text, result, max_score)

        except Exception as e:
            error = str(e)

    return render_template(
        "bjb.html",
        result=result,
        student_text=student_text,
        error=error,
        max_score=max_score
    )

# ===== ОҚУШЫЛАР =====
@app.route("/students", methods=["GET", "POST"])
def students():
    error = ""

    if request.method == "POST":
        try:
            name = request.form["student_name"]
            cls = request.form["class_name"]
            work = request.form["work_type"]
            images = request.files.getlist("images")

            for image in images:
                if image and image.filename:
                    filename = secure_filename(image.filename)
                    filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
                    image.save(filepath)

                    with open(filepath + ".txt", "w", encoding="utf-8") as f:
                        f.write(f"{name} | {cls} | {work}")

            return redirect(url_for("success"))

        except Exception as e:
            error = str(e)

    return render_template("students.html", error=error)

# ===== SUCCESS =====
@app.route("/success")
def success():
    return render_template("success.html")

# ===== МҰҒАЛІМ =====
@app.route("/teacher")
def teacher():
    files = []

    for filename in os.listdir(UPLOAD_FOLDER):
        if filename.endswith((".png", ".jpg", ".jpeg")):
            info = ""
            txt_path = os.path.join(UPLOAD_FOLDER, filename + ".txt")

            if os.path.exists(txt_path):
                with open(txt_path, "r", encoding="utf-8") as f:
                    info = f.read()

            files.append({
                "filename": filename,
                "info": info
            })

    return render_template("teacher.html", files=files)

@app.route("/uploads/<filename>")
def uploaded_file(filename):
    return send_from_directory(UPLOAD_FOLDER, filename)

# ===== ЭССЕ ҚҰРУ БЕТІ =====
@app.route("/essay-create", methods=["GET", "POST"])
def essay_create():
    essay_result = None
    error = ""

    if request.method == "POST":
        try:
            topic = request.form["topic"]
            grade_level = request.form["grade_level"]
            word_count = request.form["word_count"]

            essay_result = generate_essay(topic, grade_level, word_count)

        except Exception as e:
            error = str(e)

    return render_template("essay_create.html", essay_result=essay_result, error=error)

# ===== ЭССЕ ТЕКСЕРУ БЕТІ =====
@app.route("/essay-check", methods=["GET", "POST"])
def essay_check():
    essay_check_result = None
    error = ""

    if request.method == "POST":
        try:
            image = request.files["image"]
            topic = request.form.get("topic", "")
            grade_level = request.form.get("grade_level", "")
            word_count = request.form.get("word_count", "")

            recognized_text = extract_text_from_image(image)
            essay_check_result = check_essay_with_gemini(
                recognized_text,
                topic=topic,
                grade_level=grade_level,
                word_count=word_count
            )

        except Exception as e:
            error = str(e)

    return render_template("essay_check.html", essay_check_result=essay_check_result, error=error)

# ===== ТЕСТ ҚҰРУ БЕТІ =====
@app.route("/test-create", methods=["GET", "POST"])
def test_create():
    test_result = None
    error = ""

    if request.method == "POST":
        try:
            subject = request.form["subject"]
            grade_level = request.form["grade_level"]
            question_count = request.form["question_count"]
            difficulty = request.form["difficulty"]

            test_result = generate_test(subject, grade_level, question_count, difficulty)

        except Exception as e:
            error = str(e)

    return render_template("test_create.html", test_result=test_result, error=error)


# ===== ТЕСТ ТЕКСЕРУ БЕТІ =====
@app.route("/test-check", methods=["GET", "POST"])
def test_check():
    result = None
    extracted_test_text = ""
    extracted_student_answer_text = ""
    extracted_key = None
    error = ""

    if request.method == "POST":
        try:
            test_text = request.form.get("test_text", "").strip()
            student_answers_raw = request.form.get("student_answers", "").strip()

            test_image = request.files.get("test_image")
            student_answer_image = request.files.get("student_answer_image")

            if test_image and test_image.filename:
                extracted_test_text = extract_text_from_image(test_image)
            else:
                extracted_test_text = test_text

            if student_answer_image and student_answer_image.filename:
                extracted_student_answer_text = extract_text_from_image(student_answer_image)
                student_answers_raw = extracted_student_answer_text

            key_data = extract_answer_key_from_test_text(extracted_test_text)
            extracted_key = key_data["answer_key"]

            student_answers = parse_student_answers(student_answers_raw)
            result = check_test_answers(extracted_key, student_answers)

        except Exception as e:
            error = str(e)

    return render_template(
        "test_check.html",
        result=result,
        extracted_test_text=extracted_test_text,
        extracted_student_answer_text=extracted_student_answer_text,
        extracted_key=extracted_key,
        error=error
    )
# ===== RUN =====
if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)))